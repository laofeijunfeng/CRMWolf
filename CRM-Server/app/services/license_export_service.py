"""License 文档导出服务

导出 License 申请为 Word 文档，格式参照样例。
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import tempfile

from app.models.license_application import LicenseApplication, LicenseType
from app.utils.time import business_now


LICENSE_DOCUMENT_FONT = "黑体"


def _set_run_font(run, font_name: str = LICENSE_DOCUMENT_FONT) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_style_font(style, font_name: str = LICENSE_DOCUMENT_FONT) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_document_font(doc: Document, font_name: str = LICENSE_DOCUMENT_FONT) -> None:
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in doc.styles:
            _set_style_font(doc.styles[style_name], font_name)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font_name)


def export_license_document(application: LicenseApplication) -> str:
    """
    导出 License 文档（参照样例格式）

    Args:
        application: License 申请记录

    Returns:
        str: 导出文件的临时路径

    文件名格式：私有化部署License-{客户名称}_{当前日期}.docx
    """
    doc = Document()
    _set_document_font(doc)

    # 标题
    doc.add_heading('Apifox私有化授权文件', 0)

    # 企业信息
    doc.add_paragraph(f"企业名称: {application.customer.account_name}")
    doc.add_paragraph(f"到期时间: {application.expiry_date.strftime('%Y-%m-%d')}")

    # 授权人数来自本次 License 申请，服务器地址来自部署信息。
    doc.add_paragraph(f"授权人数: {application.authorized_users}")
    if application.deployment_info:
        doc.add_paragraph(f"服务器: {application.deployment_info.server_address}")
    else:
        doc.add_paragraph("服务器: 未配置")

    doc.add_paragraph(f"支持模块: {application.supported_modules or '未填写'}")

    # License 类型标题
    license_type_text = '试用' if application.license_type == LicenseType.TRIAL else '正式'
    authorized_users = application.authorized_users
    doc.add_paragraph()
    doc.add_paragraph(f"{license_type_text} License（{authorized_users}人）")

    # 服务端 License
    doc.add_paragraph()
    doc.add_paragraph("服务端 License:")
    if application.server_license_code:
        # 分行显示长文本
        for line in application.server_license_code.split('\n'):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("未填写")

    # 客户端 License
    doc.add_paragraph()
    doc.add_paragraph("客户端 License:")
    if application.client_license_code:
        # 分行显示长文本
        for line in application.client_license_code.split('\n'):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("未填写")

    # 文件名：私有化部署License-{客户名称}_{当前日期}.docx
    current_date = business_now().strftime('%Y%m%d')
    safe_customer_name = application.customer.account_name.translate(str.maketrans('', '', '\\/:*?"<>|\r\n'))
    file_name = f"私有化部署License-{safe_customer_name}_{current_date}.docx"
    file_path = tempfile.mktemp(suffix='.docx', prefix=file_name)
    _set_document_font(doc)
    doc.save(file_path)

    return file_path
